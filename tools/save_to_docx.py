
from docx import Document
from docx.shared import Pt
import os

def save_plan_to_docx(plan_data, output):
    """
    Converts a Test Plan dictionary to a DOCX file.
    plan_data: dict
    output: file path string OR file-like object (BytesIO)
    """
    doc = Document()
    
    # Title
    title = plan_data.get('plan_title', 'Test Plan')
    doc.add_heading(title, 0)
    
    # Sections
    sections = plan_data.get('sections', [])
    if not sections:
        doc.add_paragraph("No sections generated.")
        
    for section in sections:
        # Heading
        heading = section.get('heading', 'Section')
        doc.add_heading(heading, level=1)
        
        # Content
        content = section.get('content', '')
        
        # Simple markdown cleanup
        paragraphs = content.split('\n')
        for p_text in paragraphs:
            p_text = p_text.strip()
            if p_text:
                if p_text.startswith('* ') or p_text.startswith('- '):
                    p = doc.add_paragraph(p_text[2:], style='List Bullet')
                elif p_text.startswith('1. '): 
                     p = doc.add_paragraph(p_text, style='List Number')
                else:
                    p = doc.add_paragraph(p_text)
                    
    try:
        doc.save(output)
        return True
    except Exception as e:
        print(f"Error saving DOCX: {e}")
        return False
